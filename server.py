#!/usr/bin/env python3
"""Syllabot Micro-Server — serves static files + /upload and /generate endpoints."""

import http.server
import json
import os
import re
import tempfile
import urllib.request
import urllib.error
from html.parser import HTMLParser

PORT = 8000

# ── OpenStax URLs ──────────────────────────────────────────────────────────
OPENSTAX_BASE = "https://openstax.org/books/principles-economics-3e/pages/"
OPENSTAX_CHAPTERS = {
    1: "1-introduction",
    2: "2-introduction",
    3: "3-introduction",
    4: "4-introduction",
    5: "5-introduction",
    6: "6-introduction",
    7: "7-introduction",
    8: "8-introduction",
    9: "9-introduction",
    10: "10-introduction",
    11: "11-introduction",
    12: "12-introduction",
    13: "13-introduction",
    14: "14-introduction",
    15: "15-introduction",
}

# ── Question Templates ─────────────────────────────────────────────────────
QUESTION_TEMPLATES = {
    "definition": [
        "Which of the following best defines {term}?",
        "The concept of {term} refers to:",
        "{term} is best described as:",
    ],
    "application": [
        "Which scenario best illustrates {term}?",
        "A real-world example of {term} would be:",
    ],
    "comparison": [
        "What is the key difference between {term_a} and {term_b}?"
    ],
}

# ── HTML text extractor ────────────────────────────────────────────────────
class TextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "nav", "header", "footer"):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ("script", "style", "nav", "header", "footer"):
            self._skip = False

    def handle_data(self, data):
        if not self._skip:
            t = data.strip()
            if t:
                self.parts.append(t)

    def get_text(self):
        return "\n".join(self.parts)


# ── Content Generation Helpers ─────────────────────────────────────────────
def scrape_topic_info(topic, reading):
    """Fetch OpenStax content for a topic."""
    ch = None
    if reading:
        m = re.search(r"(\d+)", reading)
        if m:
            ch = int(m.group(1))
    scraped = ""
    if ch and ch in OPENSTAX_CHAPTERS:
        url = OPENSTAX_BASE + OPENSTAX_CHAPTERS[ch]
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Syllabot/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                html = resp.read().decode("utf-8", errors="replace")
            ext = TextExtractor()
            ext.feed(html)
            scraped = ext.get_text()[:5000]
        except Exception:
            scraped = ""
    keywords = [w.strip() for w in re.split(r"[,&]| and ", topic) if w.strip()]
    return {"topic": topic, "keywords": keywords, "reading": reading or "", "scraped_content": scraped}


def extract_key_terms(info):
    terms = []
    text = info["scraped_content"]
    # capitalized multi-word phrases
    for m in re.finditer(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b", text):
        t = m.group(1)
        if len(t) > 3 and t not in terms:
            terms.append(t)
    # topic parts
    for kw in info["keywords"]:
        if len(kw) > 3 and kw not in terms:
            terms.append(kw)
    return terms[:15]


def _find_definition(term, text):
    patterns = [
        re.compile(re.escape(term) + r"\s+(?:is|are|refers? to|means?|describes?)\s+(.+?\.)", re.I),
        re.compile(re.escape(term) + r"[:\-–]\s*(.+?\.)", re.I),
    ]
    for p in patterns:
        m = p.search(text)
        if m:
            return m.group(1).strip()
    # fallback: sentence containing term
    for sent in re.split(r"(?<=[.!?])\s+", text):
        if term.lower() in sent.lower():
            return sent.strip()
    return f"A key concept in economics related to {term}."


def generate_mcq(info, terms, count=15):
    questions = []
    text = info["scraped_content"]
    for term in terms:
        defn = _find_definition(term, text)
        # definition question
        tmpl = QUESTION_TEMPLATES["definition"][len(questions) % 3]
        questions.append({
            "q": tmpl.format(term=term),
            "opts": [
                defn,
                f"An unrelated concept from a different field of study",
                f"The opposite of what {term} actually describes",
                f"A concept that applies only in theory, not in real-world scenarios",
            ],
            "correct": 0,
            "explain": f"{term}: {defn}",
        })
        # application question
        if len(questions) < count:
            tmpl2 = QUESTION_TEMPLATES["application"][len(questions) % 2]
            questions.append({
                "q": tmpl2.format(term=term),
                "opts": [
                    f"A situation demonstrating {term.lower()} in everyday markets",
                    f"An unrelated concept from a different field of study",
                    f"The opposite of what {term} actually describes",
                    f"A concept that applies only in theory, not in real-world scenarios",
                ],
                "correct": 0,
                "explain": f"This illustrates {term} because {defn}",
            })
        if len(questions) >= count:
            break
    return questions[:count]


def generate_flashcards(info, terms, count=12):
    cards = []
    text = info["scraped_content"]
    for term in terms:
        defn = _find_definition(term, text)
        cards.append({"front": f"What is {term}?", "back": defn})
        if len(cards) < count:
            cards.append({"front": f"Why is {term} important?", "back": f"{term} matters because {defn}"})
        if len(cards) >= count:
            break
    return cards[:count]


def generate_written(info, terms, count=3):
    questions = []
    if len(terms) >= 1:
        t = terms[0]
        questions.append({
            "q": f"Explain {t} and why it matters in economics.",
            "keywords": [t.lower()] + [k.lower() for k in info["keywords"][:3]],
            "idealAnswer": f"{t} is a fundamental concept. {_find_definition(t, info['scraped_content'])}",
        })
    if len(terms) >= 3:
        a, b = terms[0], terms[2]
        questions.append({
            "q": f"Compare and contrast {a} and {b}.",
            "keywords": [a.lower(), b.lower(), "difference", "similar"],
            "idealAnswer": f"{a} and {b} are related but distinct concepts in economics.",
        })
    if len(terms) >= 2:
        t = terms[1]
        questions.append({
            "q": f"Using a real-world example, explain {t}.",
            "keywords": [t.lower(), "example", "real-world"],
            "idealAnswer": f"A real-world example of {t}: {_find_definition(t, info['scraped_content'])}",
        })
    return questions[:count]


def generate_tutor_html(info, terms):
    topic = info["topic"]
    text = info["scraped_content"]
    html = f"<div class='tutor-section'><h3>{topic}</h3>"
    if text:
        paragraphs = [p.strip() for p in text.split("\n") if len(p.strip()) > 40][:8]
        for p in paragraphs:
            # highlight key terms
            for t in terms[:5]:
                p = re.sub(re.escape(t), f"<span class='key-term'>{t}</span>", p, flags=re.I)
            html += f"<p>{p}</p>"
    if terms:
        html += "<div class='example'><strong>Key Concepts:</strong><ul>"
        for t in terms[:8]:
            html += f"<li><span class='key-term'>{t}</span>: {_find_definition(t, text)}</li>"
        html += "</ul></div>"
    html += "</div>"
    return html


# ── Multipart Parser ──────────────────────────────────────────────────────
def parse_multipart(body, content_type):
    boundary = None
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            boundary = part[9:].strip('"')
    if not boundary:
        return None, None
    parts = body.split(("--" + boundary).encode())
    for part in parts:
        if b"Content-Disposition" in part and b"filename=" in part:
            header_end = part.find(b"\r\n\r\n")
            if header_end == -1:
                continue
            header = part[:header_end].decode("utf-8", errors="replace")
            file_data = part[header_end + 4:]
            if file_data.endswith(b"\r\n"):
                file_data = file_data[:-2]
            fn_match = re.search(r'filename="([^"]+)"', header)
            filename = fn_match.group(1) if fn_match else "upload"
            return filename, file_data
    return None, None


# ── Request Handler ────────────────────────────────────────────────────────
class Handler(http.server.SimpleHTTPRequestHandler):

    def end_headers(self):
        # Disable caching so browser always gets latest index.html
        self.send_header("Cache-Control", "no-cache, no-store, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        super().end_headers()

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, GET, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors()
        self.end_headers()

    def _json_response(self, data, code=200):
        body = json.dumps(data).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self._cors()
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self):
        if self.path == "/upload":
            self._handle_upload()
        elif self.path == "/generate":
            self._handle_generate()
        else:
            self._json_response({"error": "Not found"}, 404)

    # ── /upload ────────────────────────────────────────────────────────
    def _handle_upload(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(length)
            ct = self.headers.get("Content-Type", "")
            filename, file_data = parse_multipart(body, ct)
            if not filename or not file_data:
                return self._json_response({"error": "No file found in upload"}, 400)

            ext = os.path.splitext(filename)[1].lower()
            tmp = None
            try:
                if ext == ".pdf":
                    import pdfplumber
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                    tmp.write(file_data)
                    tmp.close()
                    with pdfplumber.open(tmp.name) as pdf:
                        text = "\n".join(p.extract_text() or "" for p in pdf.pages)
                    self._json_response({"type": "text", "text": text, "filename": filename})

                elif ext == ".docx":
                    import docx
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp.write(file_data)
                    tmp.close()
                    doc = docx.Document(tmp.name)
                    parts = []
                    # Extract paragraph text
                    for p in doc.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)
                    # Extract table text (critical for grade weights, schedules, etc.)
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            # Deduplicate adjacent identical cells (merged cells repeat)
                            deduped = []
                            for c in cells:
                                if not deduped or c != deduped[-1]:
                                    deduped.append(c)
                            line = "\t".join(deduped)
                            if line.strip():
                                parts.append(line)
                    text = "\n".join(parts)
                    self._json_response({"type": "text", "text": text, "filename": filename})

                elif ext == ".txt":
                    text = file_data.decode("utf-8", errors="replace")
                    self._json_response({"type": "text", "text": text, "filename": filename})

                elif ext == ".json":
                    data = json.loads(file_data.decode("utf-8"))
                    self._json_response({"type": "json", "data": data, "filename": filename})

                else:
                    self._json_response({"error": f"Unsupported file type: {ext}"}, 400)
            finally:
                if tmp and os.path.exists(tmp.name):
                    os.unlink(tmp.name)

        except Exception as e:
            self._json_response({"error": str(e)}, 500)

    # ── /generate ──────────────────────────────────────────────────────
    def _handle_generate(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length).decode())
            modules = body.get("modules", [])
            types = body.get("types", {})
            result = {}
            if types.get("quiz"):
                result["quiz"] = {}
            if types.get("written"):
                result["written"] = {}
            if types.get("flashcards"):
                result["flashcards"] = {}

            for mod in modules:
                num = mod.get("num", 0)
                topic = mod.get("topic", "")
                reading = mod.get("reading", "")
                key = f"Module {num}: {topic}"

                info = scrape_topic_info(topic, reading)
                terms = extract_key_terms(info)

                if "quiz" in result:
                    result["quiz"][key] = generate_mcq(info, terms)
                if "written" in result:
                    result["written"][key] = generate_written(info, terms)
                if "flashcards" in result:
                    result["flashcards"][key] = generate_flashcards(info, terms)
                # always generate tutor content
                if "tutor" not in result:
                    result["tutor"] = {}
                result["tutor"][key] = generate_tutor_html(info, terms)

            self._json_response(result)
        except Exception as e:
            self._json_response({"error": str(e)}, 500)


if __name__ == "__main__":
    print(f"Syllabot server running at http://localhost:{PORT}")
    server = http.server.HTTPServer(("", PORT), Handler)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nServer stopped.")
        server.server_close()
